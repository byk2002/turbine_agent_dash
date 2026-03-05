import base64
import os
import math
import pickle
import jieba
import json
import docx
from unstructured.partition.pdf import partition_pdf
from unstructured.documents.elements import Table, Image as UnstructuredImage
# 引入必要的 LangChain 消息类
from langchain_core.messages import HumanMessage, SystemMessage
import ast  # 用于解析元数据中的列表字符串
from dataclasses import dataclass, asdict
import logging
from langchain_deepseek import ChatDeepSeek
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("langdetect").setLevel(logging.ERROR)
from typing import List, Dict, Any, Optional
from pathlib import Path
from langchain_community.document_loaders import TextLoader, PyPDFLoader, UnstructuredWordDocumentLoader
from langchain_community.retrievers import BM25Retriever
from langchain_classic.retrievers import EnsembleRetriever
from langchain_chroma import Chroma
import hashlib
from langchain_huggingface import HuggingFaceEmbeddings
from transformers import AutoTokenizer, AutoModel
from sentence_transformers import SentenceTransformer
from sentence_transformers import CrossEncoder
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_core.output_parsers import StrOutputParser

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_community.chat_message_histories import SQLChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_core.callbacks import CallbackManagerForRetrieverRun
import numpy as np
import shutil
from collections import defaultdict
from langchain_community.vectorstores.utils import filter_complex_metadata
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')  # 增加时间戳
logger = logging.getLogger(__name__)
import pytesseract
tesseract_dir = r"C:\Program Files\Tesseract-OCR"
tessdata_dir = os.path.join(tesseract_dir, "tessdata")
pytesseract.pytesseract.tesseract_cmd = os.path.join(tesseract_dir, "tesseract.exe")
# 假设你解压到了这个位置，请修改为实际路径\
os.environ["PATH"] += os.pathsep + tesseract_dir
os.environ["TESSDATA_PREFIX"] = tessdata_dir

try:
    from pix2text import Pix2Text
    # 初始化 P2T 模型 (建议在类初始化时加载，避免重复加载)
    p2t = Pix2Text.from_config()
except ImportError:
    p2t = None


class UnstructuredPDFParser:

    def __init__(self, base_path: Path):
        # 图片保存的基础路径
        self.output_dir = base_path / "assets"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def extract(self, file_path: str) -> List[Document]:
        file_path_obj = Path(file_path)
        logger.info(f"正在使用 Unstructured 解析 PDF (hi_res 模式): {file_path_obj.name} ... 这可能需要一些时间")

        try:
            # 【修改点 1】将 "Table" 加入图片提取列表
            # infer_table_structure=False (既然当图片看，就不需要浪费时间解析 HTML 结构了，除非需要 OCR 文本辅助检索)
            # 建议保留 infer_table_structure=True 以获得更好的 OCR 文本用于 Search 索引，但忽略其 HTML 输出
            elements = partition_pdf(
                filename=file_path,
                strategy="hi_res",
                extract_images_in_pdf=True,
                extract_image_block_types=["Image", "Table"],  # 【关键】把表格当图片切出来
                extract_image_block_output_dir=str(self.output_dir),
                infer_table_structure=True,  # 依然开启，为了获取 element.text 用于向量检索
                chunking_strategy="by_title",  # 使用 Unstructured 自带的语义分块作为第一层切割
                max_characters=5000,  # 限制初始块大小
                new_after_n_chars=300,
                combine_text_under_n_chars=1000,
                languages=["eng", "chi_sim"],
            )
        except Exception as e:
            logger.error(f"Unstructured 解析失败: {e}", exc_info=True)
            return []

        documents = []
        for el in elements:
            # 获取基本元数据
            page_num = el.metadata.page_number - 1 if el.metadata.page_number else 0
            file_name = file_path_obj.name

            # 准备 Document 的 metadata
            metadata = {
                "source": str(file_path_obj),
                "file_name": file_name,
                "page": page_num,
                "image_paths": []  # 初始化为空列表
            }

            # 【修改点 2】处理表格和图片元素
            if isinstance(el, Table) or isinstance(el, UnstructuredImage):
                # 尝试获取图片路径
                if el.metadata.image_path:
                    img_path = str(Path(el.metadata.image_path).resolve())
                    # 将图片路径存入 list (兼容原有逻辑)
                    metadata["image_paths"].append(img_path)
                    metadata["is_image_asset"] = True  # 标记这是个视觉资产
                # 决定 page_content
                # 对于表格，我们使用 OCR 识别到的文本作为 content，以便 BM25/Vector 能检索到它
                # 检索到后，LLM 会看 metadata 里的图片
                if el.text and el.text.strip():
                    content = el.text
                else:
                    # 如果 OCR 没识别出字，给个占位符，防止空内容被丢弃
                    content = f"[Visual Data] Type: {type(el).__name__}"
                # 标记类型，防止被 Splitter 切碎
                metadata["do_not_split"] = True
                documents.append(Document(page_content=content, metadata=metadata))
            # 【修改点 3】处理普通文本
            else:
                if el.text and el.text.strip():
                    # 普通文本，标记为 False，允许后续被 RecursiveSplitter 进一步处理
                    metadata["is_image_asset"] = False
                    metadata["do_not_split"] = False
                    documents.append(Document(page_content=el.text, metadata=metadata))

        logger.info(f"Unstructured 解析完成，生成 {len(documents)} 个原始文档片段。")
        return documents


class ChineseBM25Retriever(BM25Retriever):
    """针对中文优化的 BM25 检索器"""

    @classmethod
    def from_documents(cls, documents, **kwargs):
        # 对文档内容进行中文分词预处理
        processed_texts = []
        for doc in documents:
            # 使用 jieba 分词
            words = jieba.lcut(doc.page_content)
            processed_texts.append(" ".join(words))

        # 创建临时文档用于 BM25
        processed_docs = [
            Document(page_content=text, metadata=doc.metadata)
            for text, doc in zip(processed_texts, documents)
        ]
        return super().from_documents(processed_docs, **kwargs)

    def _get_relevant_documents(
        self,
        query: str,
        *,
        run_manager: CallbackManagerForRetrieverRun = None
    ) -> List[Document]:
        # 对查询也进行分词
        query_words = " ".join(jieba.lcut(query))
        # 将 run_manager 传递给父类
        return super()._get_relevant_documents(query_words, run_manager=run_manager)


# 【修改点 1】添加 @dataclass 装饰器
@dataclass
class ChunkData:
    """存储单个 chunk 的数据结构"""
    content: str
    metadata: Dict[str, Any]
    chunk_id: str


class ChunkStore:
    """
    Chunks 持久化存储管理器
    支持两种存储格式：
    1. JSON（可读性好，适合调试）
    2. Pickle（性能更好，适合大规模数据）
    """

    def __init__(self, store_path: Path, use_pickle: bool = True):
        self.store_path = store_path
        self.use_pickle = use_pickle
        self.store_path.mkdir(parents=True, exist_ok=True)

        # 存储文件路径
        self.chunks_file = self.store_path / ("chunks.pkl" if use_pickle else "chunks.json")
        self.index_file = self.store_path / "chunk_index.json"

        # 内存缓存
        self._chunks: Dict[str, ChunkData] = {}  # chunk_id -> ChunkData
        self._file_chunks: Dict[str, List[str]] = {}  # file_path -> [chunk_ids]
        # 【新增】辅助字典，用于不区分大小写的路径查找
        self._file_chunks_lower: Dict[str, str] = {}

        self._load()

    def _load(self):
        """从磁盘加载已存储的 chunks"""
        try:
            # 加载索引
            if self.index_file.exists():
                with open(self.index_file, 'r', encoding='utf-8') as f:
                    self._file_chunks = json.load(f)
                # 【新增】构建小写映射表
                self._file_chunks_lower = {k.lower(): k for k in self._file_chunks.keys()}
                logger.info(f"加载了 {len(self._file_chunks)} 个文件的 chunk 索引")

            # 加载 chunks 数据
            if self.chunks_file.exists():
                if self.use_pickle:
                    with open(self.chunks_file, 'rb') as f:
                        self._chunks = pickle.load(f)
                else:
                    with open(self.chunks_file, 'r', encoding='utf-8') as f:
                        raw_data = json.load(f)
                        self._chunks = {
                            k: ChunkData(**v) for k, v in raw_data.items()
                        }
                logger.info(f"加载了 {len(self._chunks)} 个 chunks")
        except Exception as e:
            logger.error(f"加载 chunk store 失败: {e}")
            self._chunks = {}
            self._file_chunks = {}
            self._file_chunks_lower = {}

    def _save(self):
        """保存 chunks 到磁盘 (增加原子写入保护)"""
        try:
            # 保存索引
            with open(self.index_file, 'w', encoding='utf-8') as f:
                json.dump(self._file_chunks, f, ensure_ascii=False, indent=2)

            # 保存 chunks 数据
            if self.use_pickle:
                temp_file = self.chunks_file.with_suffix('.tmp')
                with open(temp_file, 'wb') as f:
                    pickle.dump(self._chunks, f)
                if self.chunks_file.exists():
                    os.remove(self.chunks_file)
                os.rename(temp_file, self.chunks_file)
            else:
                with open(self.chunks_file, 'w', encoding='utf-8') as f:
                    json.dump(
                        {k: asdict(v) for k, v in self._chunks.items()},
                        f, ensure_ascii=False, indent=2
                    )
            logger.info(f"保存了 {len(self._chunks)} 个 chunks 到磁盘")
        except Exception as e:
            logger.error(f"保存 chunk store 失败: {e}")

    def add_chunks(self, file_path: str, chunks: List[Document]) -> List[str]:
        """
        添加文件的 chunks
        返回 chunk_ids 列表
        """
        chunk_ids = []
        for chunk in chunks:
            chunk_id = chunk.metadata.get('chunk_id')
            if not chunk_id:
                # 生成 chunk_id
                content_hash = hashlib.sha256(chunk.page_content.encode()).hexdigest()[:16]
                chunk_id = f"{hashlib.md5(file_path.encode()).hexdigest()[:8]}_{content_hash}"

            chunk_data = ChunkData(
                content=chunk.page_content,
                metadata=chunk.metadata,
                chunk_id=chunk_id
            )
            self._chunks[chunk_id] = chunk_data
            chunk_ids.append(chunk_id)

        self._file_chunks[file_path] = chunk_ids
        # 【新增】更新小写映射
        self._file_chunks_lower[file_path.lower()] = file_path
        self._save()
        return chunk_ids

    def get_chunks_for_file(self, file_path: str) -> List[Document]:
        """获取指定文件的所有 chunks (支持大小写模糊匹配)"""
        # 1. 尝试精确匹配
        chunk_ids = self._file_chunks.get(file_path)

        # 2. 如果失败，尝试小写模糊匹配 (解决 Windows 路径不一致问题)
        if chunk_ids is None:
            real_key = self._file_chunks_lower.get(file_path.lower())
            if real_key:
                chunk_ids = self._file_chunks.get(real_key, [])
            else:
                chunk_ids = []

        documents = []
        for chunk_id in chunk_ids:
            if chunk_id in self._chunks:
                chunk_data = self._chunks[chunk_id]
                documents.append(Document(
                    page_content=chunk_data.content,
                    metadata=chunk_data.metadata
                ))
        return documents

    def has_file(self, file_path: str) -> bool:
        """检查文件是否已存储 (支持模糊匹配)"""
        if file_path in self._file_chunks:
            return True
        return file_path.lower() in self._file_chunks_lower

    def delete_file(self, file_path: str) -> List[str]:
        """删除文件的所有 chunks，返回被删除的 chunk_ids"""
        # 尝试获取真实 Key
        real_key = file_path
        if real_key not in self._file_chunks:
            real_key = self._file_chunks_lower.get(file_path.lower())

        if not real_key:
            return []

        chunk_ids = self._file_chunks.pop(real_key, [])
        # 清理辅助映射
        if real_key.lower() in self._file_chunks_lower:
            del self._file_chunks_lower[real_key.lower()]

        for chunk_id in chunk_ids:
            self._chunks.pop(chunk_id, None)
        self._save()
        return chunk_ids

    # ... 其他方法 (get_all_chunks, get_chunk_ids_for_file, get_stats) 保持不变 ...
    def get_all_chunks(self) -> List[Document]:
        return [Document(page_content=cd.content, metadata=cd.metadata) for cd in self._chunks.values()]

    def get_chunk_ids_for_file(self, file_path: str) -> List[str]:
        return self._file_chunks.get(file_path, [])

    def get_stats(self) -> Dict[str, int]:
        return {
            "total_files": len(self._file_chunks),
            "total_chunks": len(self._chunks),
            "store_size_bytes": self.chunks_file.stat().st_size if self.chunks_file.exists() else 0
        }

class MultiDocumentKnowledgeBase:
    def __init__(self, kb_path: str,deepseek_api_key: str,llm_instance=None):
        self.kb_path = Path(kb_path).resolve()
        self.kb_path.mkdir(parents=True, exist_ok=True)
        self._deepseek_api_key = deepseek_api_key

        self.collection_name = "multi_doc_kb_collection_optimized"
        self.chroma_persist_path = str(self.kb_path / "chroma_db")
        Path(self.chroma_persist_path).mkdir(parents=True, exist_ok=True)
        logger.info(f"Chroma 存储路径: {self.chroma_persist_path}")

        self.chat_history_db_path = str(self.kb_path / "chat_histories.db")
        Path(self.chat_history_db_path).parent.mkdir(parents=True, exist_ok=True)
        logger.info(f"聊天历史数据库路径: {self.chat_history_db_path}")

        #logger.info("正在加载 HuggingFace 嵌入模型:my_sentence_transformer_model")
        '''self.embeddings = HuggingFaceEmbeddings(
            model_name="output/my_sentence_transformer_model",  # Change this to your local model directory
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )'''
        embedding_model_name = "./models/Qwen3-Embedding-0.6B"
        logger.info(f"正在从本地加载模型: {embedding_model_name}")
        self.embeddings = HuggingFaceEmbeddings(
            model_name=embedding_model_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        logger.info("HuggingFace 嵌入模型加载完成。")

        reranker_model_name = "./models/bge-reranker-large"
        # 或者直接写：reranker_model_name = "./models/bge-reranker-large"
        logger.info(f"正在从本地加载重新排名模型: {reranker_model_name}...")
        # CrossEncoder 支持直接传入本地路径
        self.reranker = CrossEncoder(reranker_model_name, device='cpu')
        logger.info("重新排名模型加载完成。")


        if llm_instance:
            self.llm = llm_instance
        else:
                # 默认回退逻辑，或者连接您的微调模型 endpoint
            self.llm = ChatDeepSeek(
                api_key=self._deepseek_api_key,  # 本地模型通常不需要真实 Key
                # 您的微调模型服务地址
                model_name="deepseek-reasoner",  # 您的模型名称
                temperature=0.2  # 专业领域建议低温度
                )

        logger.info(f"正在加载分词器: {embedding_model_name} 以实现 Token 精确切割...")
        try:
            # 加载与 Embedding 模型匹配的分词器
            self.tokenizer = AutoTokenizer.from_pretrained(embedding_model_name, trust_remote_code=True)

            # 使用 from_huggingface_tokenizer 创建分割器
            # chunk_size 现在代表 Token 数量，而不是字符数
            # 512 token 大约等于 800-1000 个汉字，适合大多数 Embedding 模型
            self.text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
                tokenizer=self.tokenizer,
                chunk_size=400,  # 限制每个块 500 tokens
                chunk_overlap=100,  # 重叠 100 tokens
                separators=["\n\n", "\n", "。", "！", "？", ".", " ", ""],  # 优先按段落和句子切
                strip_whitespace=True
            )
            logger.info(
                "文本分割器已更新为 Token 感知模式 (RecursiveCharacterTextSplitter.from_huggingface_tokenizer)。")
        except Exception as e:
            logger.error(f"加载分词器失败，回退到字符长度切割: {e}")
            # 回退逻辑
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=100,
                length_function=len
            )

        logger.info("文本分割器已初始化为 RecursiveCharacterTextSplitter。")

        self.vectorstore: Optional[Chroma] = None
        self.documents: List[Document] = []
        self.doc_metadata: Dict[str, Any] = {}
        self._file_hashes: Dict[str, str] = {}

        self._bm25_retriever: Optional[BM25Retriever] = None
        self._chroma_retriever: Optional[Any] = None
        self.ensemble_retriever: Optional[EnsembleRetriever] = None

        logger.info("聊天历史将通过 SQLChatMessageHistory 管理。")

        # 【修改点 2】必须先初始化 chunk_store，因为 _load_knowledge_base 会用到它
        self.chunk_store = ChunkStore(
            self.kb_path / "chunk_store",
            use_pickle=False  # 使用 pickle 获得更好性能
        )
        # 【新增】 初始化高级 PDF 解析器
        self.pdf_parser = UnstructuredPDFParser(self.kb_path)
        # 加载知识库必须在 chunk_store 初始化之后
        self._load_knowledge_base()

    def parse_local_file(self, file_path: str) -> str:
        """
        直接解析本地文件内容返回字符串，不进行向量存储。
        用于提取参考答案、临时上下文等。
        """
        file_path_obj = Path(file_path).resolve()
        if not file_path_obj.exists():
            logger.error(f"文件未找到: {file_path}")
            return ""

        logger.info(f"正在解析临时参考文件: {file_path_obj.name}")
        content = ""
        try:
            # 根据后缀选择解析方式
            suffix = file_path_obj.suffix.lower()

            if suffix in ['.jpg', '.jpeg', '.png', '.bmp']:
                if p2t:
                    logger.info(f"正在使用 Pix2Text 解析作业图片 (含公式): {file_path_obj.name}")
                    try:
                        # recognize_text 会自动处理混合排版并输出 LaTeX
                        res = p2t.recognize_text(str(file_path_obj))
                        content = res  # 获取识别后的 Markdown/LaTeX 文本
                    except Exception as e:
                        logger.error(f"Pix2Text 识别失败: {e}")
                else:
                    logger.warning("未安装 Pix2Text，将回退到普通 OCR")
                    # ... 原有的 Tesseract 逻辑 ...

                # --- 针对 PDF，先转图再识别 (比直接解析 PDF 文本层更准) ---
            elif suffix == '.pdf':
                # 如果是纯文本 PDF，可以用原来的逻辑
                # 但如果是扫描版/手写作业 PDF，建议转图片后用 P2T
                if p2t:
                    logger.info(f"正在将 PDF 转换为图片并使用 Pix2Text 解析: {file_path_obj.name}")
                    try:
                        from pdf2image import convert_from_path
                        images = convert_from_path(str(file_path_obj))
                        page_contents = []
                        for i, img in enumerate(images):
                            # P2T 支持直接传入 PIL Image
                            res = p2t.recognize_text(img)
                            page_contents.append(f"--- Page {i + 1} ---\n{res}")
                        content = "\n\n".join(page_contents)
                    except Exception as e:
                        logger.error(f"PDF 公式识别失败: {e}, 回退到 Unstructured")
                        # 回退到原来的 self.pdf_parser.extract
                        raw_docs = self.pdf_parser.extract(str(file_path_obj))
                        content = "\n\n".join([doc.page_content for doc in raw_docs])
                else:
                    # 原有逻辑
                    raw_docs = self.pdf_parser.extract(str(file_path_obj))
                    content = "\n\n".join([doc.page_content for doc in raw_docs])

            elif suffix == '.docx':
                print(f"DEBUG: 正在尝试使用 python-docx 解析: {file_path_obj}")  # 添加这行
                try:
                    # 优先使用 python-docx 直接解析，更稳定
                    doc = docx.Document(str(file_path_obj))
                    content = "\n".join([para.text for para in doc.paragraphs])
                    print(f"DEBUG: python-docx 解析成功，提取字符数: {len(content)}")  # 添加这行
                except Exception as e:
                    # 回退到 Unstructured
                    print(f"DEBUG: python-docx 失败，原因: {e}")  # 添加这行
                    logger.warning(f"python-docx 解析失败，尝试使用 Unstructured: {e}")
                    loader = UnstructuredWordDocumentLoader(str(file_path_obj))
                    docs = loader.load()
                    content = "\n\n".join([doc.page_content for doc in docs])
            else:
                # 默认尝试作为纯文本读取
                loader = TextLoader(str(file_path_obj), encoding='utf-8')
                docs = loader.load()
                content = "\n\n".join([doc.page_content for doc in docs])

            logger.info(f"文件解析成功，长度: {len(content)} 字符")
            return content

        except Exception as e:
            logger.error(f"解析文件失败: {e}", exc_info=True)
            return f"Error parsing file: {str(e)}"

    def _process_unstructured_documents(self, raw_documents: List[Document]) -> List[Document]:
        """
        改进版：智能选择是否需要二次切割
        """
        final_chunks = []

        for doc in raw_documents:
            # ✅ 保护的内容（表格/图片）完全不切
            if doc.metadata.get("do_not_split", False):
                final_chunks.append(doc)
                continue

            # 📊 统计文本 Token 数
            doc_tokens = len(self.tokenizer.encode(doc.page_content))

            # ✅ 如果已经在合理范围内（<600 tokens），就不再切割
            if doc_tokens <= 500:
                doc.metadata["section_context"] = doc.page_content[: 100]
                final_chunks.append(doc)
                continue

            # ❌ 只有当文本过长时才进行二次切割
            sub_chunks = self.text_splitter.split_documents([doc])

            # 为子块添加上下文
            potential_title = doc.page_content.split('\n')[0][:100]
            for i, sub_chunk in enumerate(sub_chunks):
                if i > 0:
                    sub_chunk.page_content = f"【上下文:  {potential_title}】\n{sub_chunk.page_content}"
                sub_chunk.metadata["section_context"] = potential_title
                final_chunks.append(sub_chunk)

        return final_chunks

    def _normalize_path(self, file_path: Any) -> str:
        """
        统一路径格式：绝对路径 + 统一转为小写。
        解决 Windows 下 D: 和 d: 不一致导致无法命中缓存的问题。
        """
        try:
            return str(Path(file_path).resolve()).lower()
        except Exception:
            return str(file_path).lower()

    def _calculate_file_hash(self, file_path: Path) -> str:
        hasher = hashlib.md5()
        try:
            with open(file_path, 'rb') as f:
                while chunk := f.read(4096):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"计算文件哈希时出错 {file_path}: {e}")
            return ""

    def _generate_chunk_id(self, file_path_str: str, page_number: Any, chunk_content: str) -> str:
        """为每个文档片段生成一个稳定的、唯一的ID。"""
        content_hash = hashlib.sha256(chunk_content.encode('utf-8')).hexdigest()[:16]
        unique_id_str = f"{file_path_str}#{page_number}#{content_hash}"
        return hashlib.md5(unique_id_str.encode('utf-8')).hexdigest()

    def _get_or_create_chroma(self):
        """确保 Chroma 实例存在，如果不存在则尝试从持久化路径加载或创建一个新的。"""
        if self.vectorstore is None:
            try:
                if Path(self.chroma_persist_path).exists() and list(Path(self.chroma_persist_path).iterdir()):
                    self.vectorstore = Chroma(
                        collection_name=self.collection_name,
                        embedding_function=self.embeddings,
                        persist_directory=self.chroma_persist_path
                    )
                    try:
                        _ = self.vectorstore._collection.count()
                        logger.info(f"✅ 成功恢复 Chroma 集合: {self.collection_name}")
                    except Exception as e:
                        logger.warning(f"Chroma 集合计数失败，可能已损坏或为空：{e}。将尝试创建新集合。")
                        self.vectorstore = None
                else:
                    logger.info(f"Chroma 存储路径 '{self.chroma_persist_path}' 为空或不存在，将创建一个新的。")
            except Exception as e:
                logger.warning(f"无法恢复 Chroma 实例：{e}。将创建一个新的。")
                self.vectorstore = None
        return self.vectorstore

    def _load_knowledge_base(self):
        """
        优化后的知识库加载逻辑 (v2.0 稳定版)：
        1. 引入路径标准化 (_normalize_path)，解决 Windows 盘符大小写问题。
        2. 优先比对 Hash，只要 Hash 没变，绝对不重新跑 OCR。
        3. 增强 ChunkStore 读取的容错性。
        """
        try:
            # 1. 加载元数据
            metadata_file = self.kb_path / "metadata.json"
            if metadata_file.exists():
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.doc_metadata = data.get("doc_metadata", {})
                    self._file_hashes = data.get("file_hashes", {})
                logger.info(f"加载了 {len(self.doc_metadata)} 个文件元数据。")
            else:
                self.doc_metadata = {}
                self._file_hashes = {}

            # 2. 初始化 Chroma (尝试加载持久化数据)
            self._get_or_create_chroma()

            self.documents = []  # 清空内存文档列表 (用于 BM25)
            files_to_update_in_chroma = []
            files_to_remove_from_metadata = []

            # 3. 遍历元数据中的文件
            # 使用 list() 包装 keys，防止在迭代中修改字典报错
            for stored_path_str, meta in list(self.doc_metadata.items()):
                try:
                    file_path_obj = Path(stored_path_str).resolve()
                except Exception as e:
                    logger.warning(f"路径解析失败: {stored_path_str}, 将移除记录。")
                    files_to_remove_from_metadata.append(stored_path_str)
                    continue

                # --- 检查文件是否存在 ---
                if not file_path_obj.exists():
                    logger.warning(f"文件缺失: {file_path_obj.name}，将从知识库移除。")
                    files_to_remove_from_metadata.append(stored_path_str)
                    continue

                # --- 关键：标准化路径用于后续逻辑 ---
                normalized_path = self._normalize_path(file_path_obj)

                # --- 核心优化：基于 mtime 和 Hash 的多级检查 ---
                current_mtime = str(os.path.getmtime(file_path_obj))
                last_mtime = meta.get("added_at", "0")

                file_needs_reprocess = False

                # 只有时间戳变了，才值得去算 Hash
                if current_mtime != last_mtime:
                    logger.info(f"检测到时间变化: {file_path_obj.name}，正在校验内容哈希...")
                    current_hash = self._calculate_file_hash(file_path_obj)
                    stored_hash = self._file_hashes.get(stored_path_str)

                    if current_hash != stored_hash:
                        logger.info(f"⚠️ 文件内容已修改 (Hash不一致): {file_path_obj.name}，标记为更新。")
                        file_needs_reprocess = True
                    else:
                        # 内容没变，只是碰了时间 (touch)，更新元数据中的时间即可
                        logger.info(f"✅ 文件内容未变 (Hash一致): {file_path_obj.name}，仅更新时间戳。")
                        self.doc_metadata[stored_path_str]["added_at"] = current_mtime

                # --- 加载 Chunks (用于 BM25) ---
                # 【容错读取】：尝试用原始 Key 读取，如果失败，尝试用标准化 Key 读取
                chunks = self.chunk_store.get_chunks_for_file(stored_path_str)
                if not chunks:
                    # 尝试用标准化路径读取 (应对 Windows 盘符大小写变化)
                    chunks = self.chunk_store.get_chunks_for_file(normalized_path)

                # 如果依然没读到，且文件内容没变，说明是 ChunkStore 缓存文件损坏或丢了
                if not chunks and not file_needs_reprocess:
                    logger.warning(f"❌ ChunkStore 缓存缺失: {file_path_obj.name}，必须强制重新解析。")
                    file_needs_reprocess = True

                # --- 分支 A：快乐路径 (直接使用缓存) ---
                if not file_needs_reprocess:
                    # 直接把缓存的 chunks 加到内存给 BM25 用
                    # 绝对不重新 OCR，也不操作 Chroma
                    self.documents.extend(chunks)
                    continue

                # --- 分支 B：慢速路径 (重新解析) ---
                try:
                    logger.info(f"🔄 正在重新解析文件: {file_path_obj.name} ...")

                    doc_type = meta.get("doc_type", "auto")

                    # 1. 执行解析 (耗时操作)
                    if doc_type == "pdf":
                        raw_docs = self.pdf_parser.extract(str(file_path_obj))
                        chunks = self._process_unstructured_documents(raw_docs)
                    elif doc_type == "text":
                        loader = TextLoader(str(file_path_obj), encoding='utf-8')
                        chunks = self.text_splitter.split_documents(loader.load())
                    elif doc_type == "docx":
                        loader = UnstructuredWordDocumentLoader(str(file_path_obj))
                        chunks = self.text_splitter.split_documents(loader.load())
                    else:
                        # 兜底
                        loader = TextLoader(str(file_path_obj), encoding='utf-8')
                        chunks = self.text_splitter.split_documents(loader.load())

                    # 2. 重新生成 ID 并保存到 Store
                    chunk_ids = []
                    for chunk in chunks:
                        # 补全 metadata
                        page = chunk.metadata.get('page', 0)
                        # 注意：这里使用 stored_path_str 保持一致性
                        cid = self._generate_chunk_id(stored_path_str, page, chunk.page_content)

                        raw_image_paths = chunk.metadata.get("image_paths", [])
                        chunk.metadata.update({
                            "source": stored_path_str,
                            "chunk_id": cid,
                            "file_name": file_path_obj.name,
                            "image_paths": str(raw_image_paths)
                        })
                        chunk_ids.append(cid)

                    # 3. 更新 ChunkStore
                    # 建议：这里如果 stored_path_str 和 normalized_path 不一致，
                    # 为了未来稳定性，可以考虑删除旧 key，存入新 key。
                    # 但为了简单，我们还是存入 stored_path_str
                    self.chunk_store.add_chunks(stored_path_str, chunks)

                    # 4. 更新元数据
                    self.doc_metadata[stored_path_str].update({
                        "chunk_ids": chunk_ids,
                        "chunk_count": len(chunks),
                        "added_at": current_mtime
                    })
                    self._file_hashes[stored_path_str] = self._calculate_file_hash(file_path_obj)

                    # 5. 收集到列表，稍后更新 Chroma
                    files_to_update_in_chroma.append(
                        (stored_path_str, chunks, chunk_ids)
                    )

                    # 6. 加入内存 (BM25)
                    self.documents.extend(chunks)

                except Exception as e:
                    logger.error(f"处理文件失败 {stored_path_str}: {e}", exc_info=True)
                    # 如果解析失败，不应该保留错误的元数据，下次应该重试
                    files_to_remove_from_metadata.append(stored_path_str)

            # 4. 执行清理和 Chroma 更新

            # A. 移除失效文件
            kb_has_changes = False
            if files_to_remove_from_metadata or files_to_update_in_chroma:
                kb_has_changes = True

            # A. 移除失效文件
            for fp in files_to_remove_from_metadata:
                if fp in self.doc_metadata: del self.doc_metadata[fp]
                if fp in self._file_hashes: del self._file_hashes[fp]
                self._delete_file_from_vectorstore(fp)

            # B. 增量更新 Chroma
            if files_to_update_in_chroma:
                self._get_or_create_chroma()
                if self.vectorstore:
                    for fp, chunks, cids in files_to_update_in_chroma:
                        logger.info(f"正在更新 Chroma 索引: {Path(fp).name} (清理旧记录 -> 添加新记录)")
                        self._delete_file_from_vectorstore(fp)
                        if chunks:
                            self.vectorstore.add_documents(chunks, ids=cids)

            # 5. 保存更新后的元数据
            self._save_knowledge_base()
            logger.info(f"知识库加载完成。内存中文档片段数: {len(self.documents)}")

        except Exception as e:
            logger.error(f"加载知识库时发生未捕获异常: {e}", exc_info=True)
            self.documents = []
            kb_has_changes = True  # 发生异常时，安全起见标记为需要重建
        finally:
            # 【修改点】根据检测到的变更状态，决定是否使用缓存
            if 'kb_has_changes' not in locals():
                kb_has_changes = True  # 兜底

            self._initialize_retrievers(force_rebuild=kb_has_changes)

    def _migrate_old_data(self):
        """迁移旧版本数据到新的 chunk store 格式"""
        for file_path_str, meta in list(self.doc_metadata.items()):
            file_path_obj = Path(file_path_str)

            if not file_path_obj.exists():
                logger.warning(f"文件 {file_path_obj.name} 不存在，跳过迁移")
                continue

            try:
                # 根据文档类型获取加载器
                doc_type = meta.get("doc_type")
                if doc_type == "pdf":
                    loader = PyPDFLoader(str(file_path_obj))
                elif doc_type == "text":
                    loader = TextLoader(str(file_path_obj), encoding='utf-8')
                elif doc_type == "docx":
                    loader = UnstructuredWordDocumentLoader(str(file_path_obj))
                else:
                    continue

                # 加载并分割
                raw_documents = loader.load()
                chunks = self.text_splitter.split_documents(raw_documents)

                # 添加元数据
                for chunk in chunks:
                    page_number = chunk.metadata.get('page', -1) + 1 if 'page' in chunk.metadata else 'N/A'
                    chunk.metadata.update({
                        "source": file_path_str,
                        "doc_type": doc_type,
                        "file_name": file_path_obj.name,
                        "page_number": page_number
                    })

                # 存储到 chunk store
                self.chunk_store.add_chunks(file_path_str, chunks)
                self.documents.extend(chunks)

                logger.info(f"已迁移文档: {file_path_obj.name} ({len(chunks)} chunks)")

            except Exception as e:
                logger.error(f"迁移文档 {file_path_obj.name} 失败: {e}")

        logger.info("数据迁移完成")

    def _save_knowledge_base(self):
        try:
            metadata_file = self.kb_path / "metadata.json"
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump({
                    "doc_metadata": self.doc_metadata,
                    "file_hashes": self._file_hashes
                }, f, ensure_ascii=False, indent=2)
            logger.info("文档元数据和文件哈希保存成功。")

        except Exception as e:
            logger.error(f"保存知识库时出错: {e}", exc_info=True)

    def _initialize_retrievers(self, force_rebuild: bool = True):
        bm25_cache_path = self.kb_path / "bm25_retriever.pkl"
        loaded_from_cache = False
        # 1. 尝试从缓存加载 (当不需要强制重建且缓存文件存在时)
        if not force_rebuild and bm25_cache_path.exists():
            try:
                with open(bm25_cache_path, 'rb') as f:
                    self._bm25_retriever = pickle.load(f)
                logger.info(f"🚀 成功从本地缓存加载 BM25 索引: {bm25_cache_path}")
                loaded_from_cache = True
            except Exception as e:
                logger.warning(f"⚠️ 加载 BM25 缓存失败: {e}，将重新构建。")

        # 2. 如果未从缓存加载，则重新构建
        if not loaded_from_cache:
            if self.documents:
                # 【优化】使用您定义的 ChineseBM25Retriever 进行更好的中文分词
                # 如果您想用回普通的 BM25Retriever，请改回 BM25Retriever.from_documents
                try:
                    self._bm25_retriever = ChineseBM25Retriever.from_documents(self.documents)
                except NameError:
                    # 以此防备ChineseBM25Retriever未定义的情况
                    self._bm25_retriever = BM25Retriever.from_documents(self.documents)

                logger.info(f"成功初始化 BM25 Retriever, 包含 {len(self.documents)} 个文档片段。")

                # 构建完成后，保存到缓存
                try:
                    with open(bm25_cache_path, 'wb') as f:
                        pickle.dump(self._bm25_retriever, f)
                    logger.info("💾 BM25 索引已保存到本地缓存。")
                except Exception as e:
                    logger.error(f"保存 BM25 缓存失败: {e}")
            else:
                self._bm25_retriever = None
                logger.warning("没有文档片段 (self.documents 为空), BM25 Retriever 未初始化。")
                # 如果没有文档，清理可能存在的旧缓存，防止不一致
                if bm25_cache_path.exists():
                    try:
                        os.remove(bm25_cache_path)
                    except OSError:
                        pass

        self._get_or_create_chroma()
        if self.vectorstore:
            try:
                # 显式指定 k 参数，即使在初始化检索器时，也可以设置一个默认值
                self._chroma_retriever = self.vectorstore.as_retriever(search_kwargs={"k": 8})
                logger.info("成功初始化 Chroma Retriever。")
            except Exception as e:
                logger.warning(f"初始化 Chroma Retriever 失败: {e}", exc_info=True)
                self._chroma_retriever = None
        else:
            self._chroma_retriever = None
            logger.warning("Chroma 向量存储未初始化, 无法创建向量检索器。")

        if self._bm25_retriever and self._chroma_retriever:
            logger.info("混合检索准备就绪: 将在 search() 中使用 RRF 融合 BM25 和 Vector 结果。")
        elif self._bm25_retriever or self._chroma_retriever:
            logger.info("仅部分检索器可用，search() 将降级为单路检索。")
        else:
            logger.error("未能初始化任何检索器。查询功能可能无法正常工作。")

    def _rrf_fusion(self, results_list: List[List[Document]], k: int = 60) -> List[Document]:
        """
        实现 RRF (Reciprocal Rank Fusion) 算法。
        公式: score = sum(1 / (k + rank))
        k 通常取 60。
        """
        fused_scores = defaultdict(float)
        doc_map = {}

        for results in results_list:
            for rank, doc in enumerate(results):
                # 生成唯一键：优先使用 chunk_id，如果没有则使用内容哈希
                doc_id = doc.metadata.get("chunk_id")
                if not doc_id:
                    unique_str = f"{doc.metadata.get('source', '')}#{doc.metadata.get('page_number', '')}#{doc.page_content}"
                    doc_id = hashlib.md5(unique_str.encode('utf-8')).hexdigest()

                if doc_id not in doc_map:
                    doc_map[doc_id] = doc

                # RRF 核心公式
                fused_scores[doc_id] += 1.0 / (k + rank + 1)

        # 根据分数降序排序
        sorted_doc_ids = sorted(fused_scores.keys(), key=lambda x: fused_scores[x], reverse=True)

        # 返回排序后的文档列表
        return [doc_map[doc_id] for doc_id in sorted_doc_ids]

    def _delete_file_from_vectorstore(self, file_path_str: str):
        self._get_or_create_chroma()
        if not self.vectorstore:
            logger.warning("向量存储未初始化，跳过删除 Chroma 记录。")
            return

        chunk_ids_to_delete = []
        if file_path_str in self.doc_metadata:
            chunk_ids_to_delete = self.doc_metadata[file_path_str].get("chunk_ids", [])

        if chunk_ids_to_delete:
            try:
                self.vectorstore.delete(ids=chunk_ids_to_delete)
                logger.info(
                    f"从 Chroma 成功删除与 '{file_path_str}' 相关的 {len(chunk_ids_to_delete)} 条记录 (通过 chunk_ids)。")
            except Exception as e:
                logger.error(f"从 Chroma 删除文档 '{file_path_str}' (使用 chunk_ids) 时出错: {e}", exc_info=True)
        else:
            logger.info(
                f"文件 '{file_path_str}' 没有关联的 chunk_ids 或未在元数据中找到，尝试通过源文件名从 Chroma 中删除所有相关片段。")

        try:
            self.vectorstore.delete(where={"source": file_path_str})
            logger.info(f"从 Chroma 成功删除与 '{file_path_str}' 相关的记录 (通过 metadata 过滤)。")
        except Exception as e:
            logger.error(f"从 Chroma 删除文档 '{file_path_str}' (使用 metadata 过滤) 时出错: {e}", exc_info=True)

    def add_document(self, file_path: str, doc_type: str = "auto"):
        file_path_obj = Path(file_path).resolve()
        if not file_path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path_obj}")

        # 【修改点】使用标准化路径作为唯一标识 Key
        file_path_str = self._normalize_path(file_path_obj)

        # 计算哈希用于比对内容是否变化
        current_file_hash = self._calculate_file_hash(file_path_obj)

        # 检查文件是否已存在且未变化
        if (file_path_str in self.doc_metadata and
                self._file_hashes.get(file_path_str) == current_file_hash and
                self.chunk_store.has_file(file_path_str)):
            logger.info(f"文档 '{file_path_obj.name}' 已存在且未变化，跳过")
            return

        # 如果文件存在但内容变了，先清理旧数据
        if file_path_str in self.doc_metadata:
            logger.info(f"文档 '{file_path_obj.name}' 内容已更新，重新处理")
            self._delete_file_from_vectorstore(file_path_str)
            self.chunk_store.delete_file(file_path_str)  # 从 chunk store 删除
            self.documents = [doc for doc in self.documents if doc.metadata.get('source') != file_path_str]

        logger.info(f"正在添加文档: {file_path_obj}")

        try:
            # 1. 加载文档 (解析)
            if doc_type == "pdf" or (doc_type == "auto" and file_path_obj.suffix.lower() == '.pdf'):
                doc_type = "pdf"
                raw_documents = self.pdf_parser.extract(file_path_str)
                # 【修改点】针对 PDF 使用自定义的混合切割逻辑
                chunks = self._process_unstructured_documents(raw_documents)
            else:
                if doc_type == "auto":
                    if file_path_obj.suffix.lower() in ['.txt', '.md', '.rtf']:
                        doc_type = "text"
                        loader = TextLoader(str(file_path_obj), encoding='utf-8')
                    elif file_path_obj.suffix.lower() == '.docx':
                        doc_type = "docx"
                        loader = UnstructuredWordDocumentLoader(str(file_path_obj))
                    else:
                        logger.warning(f"不支持的文件类型: {file_path_obj.suffix}")
                        return
                    raw_documents = loader.load()
                chunks = self.text_splitter.split_documents(raw_documents)

            # 3. 空文档检查
            if not chunks:
                logger.warning(f"文档 {file_path_obj.name} 没有生成任何文本片段,跳过。")
                if file_path_str in self.doc_metadata:
                    del self.doc_metadata[file_path_str]
                    del self._file_hashes[file_path_str]
                    self._save_knowledge_base()
                    self._initialize_retrievers()
                return

            # 4. 遍历处理 Metadata 和生成 ID
            chunk_ids_for_this_file = []
            for i, chunk in enumerate(chunks):
                # 兼容 PDFParser 或其他 Loader 的 page 字段
                if 'page' in chunk.metadata:
                    page_number = chunk.metadata['page'] + 1
                else:
                    page_number = 'N/A'

                # 生成唯一 ID
                chunk_id = self._generate_chunk_id(file_path_str, page_number, chunk.page_content)
                chunk_ids_for_this_file.append(chunk_id)

                # 获取原始的图片路径列表
                raw_image_paths = chunk.metadata.get("image_paths", [])

                # 更新 metadata (原地修改)
                chunk.metadata.update({
                    "source": file_path_str,
                    "doc_type": doc_type,
                    "file_name": file_path_obj.name,
                    "page_number": page_number,
                    "chunk_id": chunk_id,
                    # 【修复点】Chroma 不支持列表类型元数据，必须将其转为字符串
                    "image_paths": str(raw_image_paths)
                })

            # =======================================================
            # 关键修改：以下所有存储操作必须在 for 循环外部执行
            # =======================================================

            # 5. 保存到 ChunkStore
            self.chunk_store.add_chunks(file_path_str, chunks)

            # 6. 保存到 Chroma 向量数据库
            self._get_or_create_chroma()
            if self.vectorstore is None:
                self.vectorstore = Chroma.from_documents(
                    documents=chunks,
                    embedding=self.embeddings,
                    persist_directory=self.chroma_persist_path,
                    collection_name=self.collection_name,
                    ids=chunk_ids_for_this_file
                )
                logger.info(
                    f"创建了新的 Chroma 集合 '{self.collection_name}' 并持久化到 {self.chroma_persist_path}。")
            else:
                # 此时 chunks 和 chunk_ids_for_this_file 长度完全一致
                self.vectorstore.add_documents(chunks, ids=chunk_ids_for_this_file)
                logger.info(f"向 Chroma 集合追加了 {len(chunks)} 个文档片段。")

            # 7. 更新内存文档列表 (用于 BM25)
            self.documents.extend(chunks)

            # 8. 更新全局元数据
            self.doc_metadata[file_path_str] = {
                "chunk_count": len(chunks),
                "doc_type": doc_type,
                "added_at": str(os.path.getmtime(file_path_obj)),
                "chunk_ids": chunk_ids_for_this_file
            }
            self._file_hashes[file_path_str] = current_file_hash

            # 9. 持久化元数据并重新初始化检索器
            self._save_knowledge_base()
            self._initialize_retrievers()

            logger.info(f"成功添加文档: {file_path_obj} (分割为 {len(chunks)} 个片段)")

        except Exception as e:
            logger.error(f"添加文档时发生未知错误: {e}", exc_info=True)
            # 发生异常时进行清理
            if file_path_str in self.doc_metadata:
                del self.doc_metadata[file_path_str]
                del self._file_hashes[file_path_str]
                self._save_knowledge_base()
                self._delete_file_from_vectorstore(file_path_str)
                self._initialize_retrievers()
            raise

    def add_documents_from_directory(self, directory_path: str):
        directory_path_obj = Path(directory_path).resolve()
        if not directory_path_obj.is_dir():
            raise NotADirectoryError(f"路径不是一个有效的目录: {directory_path_obj}")

        logger.info(f"开始从目录 {directory_path_obj} 添加文档...")
        processed_count = 0
        all_new_chunks_for_store: List[Document] = []
        all_new_chunk_ids_for_store: List[str] = []
        processed_file_paths_metadata_temp: Dict[str, Any] = {}
        processed_file_hashes_new: Dict[str, str] = {}
        # 注意：这里 doc_type 应该在循环内根据文件判断，或者是参数传入。
        # 你的原代码在循环外定义了 doc_type="auto"，这可能导致后续混乱，建议保留默认行为。
        default_doc_type = "auto"

        files_to_delete_from_chroma_sources = []
        files_to_delete_from_memory_paths = []

        # --- 预处理：检查哪些文件被删除了 ---
        current_known_files_in_dir = set()
        for root, _, files in os.walk(directory_path_obj):
            for file_name in files:
                file_path = Path(root) / file_name
                # 【修改点 1】使用统一的路径标准化方法，防止大小写问题
                file_path_str = self._normalize_path(file_path)
                current_known_files_in_dir.add(file_path_str)

        # 检查元数据中属于该目录但现在不存在的文件
        for file_path_str in list(self.doc_metadata.keys()):
            # 简单的判断：如果路径以目录开头，且不在当前扫描列表中
            if file_path_str.startswith(
                    self._normalize_path(directory_path_obj)) and file_path_str not in current_known_files_in_dir:
                logger.info(f"目录中不再存在文件 '{Path(file_path_str).name}', 将从知识库中移除。")
                files_to_delete_from_chroma_sources.append(file_path_str)
                files_to_delete_from_memory_paths.append(file_path_str)

        # --- 主循环：处理新增或修改的文件 ---
        for root, _, files in os.walk(directory_path_obj):
            for file_name in files:
                file_path = Path(root) / file_name

                # 【修改点 1】再次确保这里使用 normalize_path，与上方和 chunk_store 保持一致
                file_path_str = self._normalize_path(file_path)

                current_file_hash = self._calculate_file_hash(file_path)

                # 检查是否已存在且未修改
                if file_path_str in self.doc_metadata and self._file_hashes.get(file_path_str) == current_file_hash:
                    # 即使文件没变，也要确保 chunk_store 里有它 (防止之前只存了元数据没存chunks的情况)
                    if self.chunk_store.has_file(file_path_str):
                        logger.info(f"文档 '{file_path.name}' 已存在且内容未变,跳过处理。")
                        processed_file_paths_metadata_temp[file_path_str] = self.doc_metadata[file_path_str]
                        processed_file_hashes_new[file_path_str] = self._file_hashes[file_path_str]
                        continue
                    else:
                        logger.warning(f"文档 '{file_path.name}' 元数据存在但 ChunkStore 缺失，将重新处理。")

                if file_path_str in self.doc_metadata:
                    logger.info(f"文档 '{file_path.name}' 内容已更新(或缓存缺失), 正在重新处理。")
                    files_to_delete_from_chroma_sources.append(file_path_str)
                    files_to_delete_from_memory_paths.append(file_path_str)

                try:
                    # 重置 doc_type 为 auto，根据每个文件后缀判断
                    current_doc_type = default_doc_type

                    # 【逻辑保持不变】解析文档
                    if current_doc_type == "pdf" or (current_doc_type == "auto" and file_path.suffix.lower() == '.pdf'):
                        current_doc_type = "pdf"
                        raw_documents = self.pdf_parser.extract(file_path_str)  # 这里传入路径字符串
                        chunks = self._process_unstructured_documents(raw_documents)
                    else:
                        if current_doc_type == "auto":
                            if file_path.suffix.lower() in ['.txt', '.md', '.rtf']:
                                current_doc_type = "text"
                                loader = TextLoader(str(file_path), encoding='utf-8')
                            elif file_path.suffix.lower() == '.docx':
                                current_doc_type = "docx"
                                loader = UnstructuredWordDocumentLoader(str(file_path))
                            else:
                                logger.warning(f"不支持的文件类型: {file_path.suffix}")
                                continue  # 使用 continue 而不是 return，防止打断整个循环
                            raw_documents = loader.load()
                        chunks = self.text_splitter.split_documents(raw_documents)

                    if not chunks:
                        logger.warning(f"文档 {file_path.name} 没有生成任何文本片段,跳过。")
                        continue

                    chunk_ids_for_this_file = []
                    for i, chunk in enumerate(chunks):
                        if 'page' in chunk.metadata:
                            page_number = chunk.metadata['page'] + 1
                        else:
                            page_number = 'N/A'

                        chunk_id = self._generate_chunk_id(file_path_str, page_number, chunk.page_content)
                        chunk_ids_for_this_file.append(chunk_id)

                        raw_image_paths = chunk.metadata.get("image_paths", [])
                        chunk.metadata.update({
                            "source": file_path_str,
                            "doc_type": current_doc_type,
                            "file_name": file_path.name,
                            "page_number": page_number,
                            "chunk_id": chunk_id,
                            "image_paths": str(raw_image_paths)
                        })

                    # =======================================================
                    # 【修改点 2】关键修复：必须将 chunks 保存到 ChunkStore 磁盘缓存！
                    # =======================================================
                    self.chunk_store.add_chunks(file_path_str, chunks)
                    # =======================================================

                    all_new_chunks_for_store.extend(chunks)
                    all_new_chunk_ids_for_store.extend(chunk_ids_for_this_file)
                    processed_count += 1

                    # 记录元数据，稍后统一更新
                    processed_file_paths_metadata_temp[file_path_str] = {
                        "chunk_count": len(chunks),
                        "doc_type": current_doc_type,
                        "added_at": str(os.path.getmtime(file_path)),
                        "chunk_ids": chunk_ids_for_this_file
                    }
                    processed_file_hashes_new[file_path_str] = current_file_hash

                    logger.info(f"已处理文档: {file_path.name}, 缓存已保存，待添加到向量存储。")

                except Exception as e:
                    logger.error(f"处理文件 {file_path} 时出错: {e}", exc_info=True)
                    if file_path_str in self.doc_metadata:
                        files_to_delete_from_chroma_sources.append(file_path_str)
                        files_to_delete_from_memory_paths.append(file_path_str)

        # --- 后续清理和向量库更新逻辑保持不变 ---

        self._get_or_create_chroma()
        # 1. 从 Chroma 删除旧数据
        if files_to_delete_from_chroma_sources and self.vectorstore:
            for file_path_to_delete in set(files_to_delete_from_chroma_sources):
                try:
                    self.vectorstore.delete(where={"source": file_path_to_delete})
                except Exception as e:
                    logger.error(f"从 Chroma 删除出错: {e}")

        # 2. 从内存和元数据中清理
        if files_to_delete_from_memory_paths:
            self.documents = [doc for doc in self.documents if
                              doc.metadata.get('source') not in files_to_delete_from_memory_paths]
            for file_path_str in set(files_to_delete_from_memory_paths):
                if file_path_str in self.doc_metadata:
                    del self.doc_metadata[file_path_str]
                if file_path_str in self._file_hashes:
                    del self._file_hashes[file_path_str]
                # 同时也从 chunk_store 清理旧文件
                self.chunk_store.delete_file(file_path_str)

        if not all_new_chunks_for_store:
            logger.info("没有需要添加到向量存储的新文档。")
            self._save_knowledge_base()
            self._initialize_retrievers()
            return

        # 3. 批量添加到 Chroma
        BATCH_SIZE = 5024
        total_chunks = len(all_new_chunks_for_store)

        try:
            self._get_or_create_chroma()
            if self.vectorstore is None:
                logger.info(f"Chroma 向量存储首次创建，正在添加 {total_chunks} 个片段...")
                self.vectorstore = Chroma.from_documents(
                    documents=all_new_chunks_for_store,
                    embedding=self.embeddings,
                    persist_directory=self.chroma_persist_path,
                    collection_name=self.collection_name,
                    ids=all_new_chunk_ids_for_store
                )
            else:
                logger.info(f"向已存在 Chroma 集合追加文档，共 {total_chunks} 个片段。")
                for i in range(0, total_chunks, BATCH_SIZE):
                    batch_chunks = all_new_chunks_for_store[i:i + BATCH_SIZE]
                    batch_ids = all_new_chunk_ids_for_store[i:i + BATCH_SIZE]
                    self.vectorstore.add_documents(batch_chunks, ids=batch_ids)

            # 4. 更新内存中的元数据
            for file_path_str, meta in processed_file_paths_metadata_temp.items():
                self.documents.extend(
                    [doc for doc in all_new_chunks_for_store if doc.metadata.get('source') == file_path_str])
                self.doc_metadata[file_path_str] = meta
                self._file_hashes[file_path_str] = processed_file_hashes_new[file_path_str]

            self._save_knowledge_base()
            self._initialize_retrievers()

            logger.info(f"完成从目录 {directory_path_obj} 添加文档。共处理/更新 {processed_count} 个文档。")
        except Exception as e:
            logger.error(f"将文档添加到 Chroma 时出错: {e}", exc_info=True)
            raise

    def _evaluate_history_relevance(self, question: str, chat_history_list: List[Any]) -> float:
        if not chat_history_list:
            return 0.0

        history_texts = []
        for msg in chat_history_list:
            if hasattr(msg, 'content'):
                history_texts.append(msg.content)
            elif isinstance(msg, dict) and 'content' in msg:
                history_texts.append(msg['content'])

        if not history_texts:
            return 0.0

        sentence_pairs = []
        for hist_text in history_texts:
            sentence_pairs.append([question, hist_text])

        try:
            scores = self.reranker.predict(sentence_pairs)
            return np.mean(scores).item()
        except Exception as e:
            logger.warning(f"评估历史聊天相关性时出错: {e}. 返回默认低分数。", exc_info=True)
            return 0.1

    def _rewrite_queries(self, question: str, chat_history_list: List[Any], num_queries: int = 5,
                         history_relevance_threshold: float = 0.5) -> list[str]:
        use_history_for_rewrite = False
        if chat_history_list:
            history_relevance = self._evaluate_history_relevance(question, chat_history_list)
            logger.info(f"问题与历史聊天的相关性分数 (用于查询重写): {history_relevance:.4f}")
            if history_relevance >= history_relevance_threshold:
                use_history_for_rewrite = True
                logger.info("问题与历史聊天相关度高，将在查询重写中考虑历史对话。")
            else:
                logger.info("问题与历史聊天相关度低，查询重写将仅基于当前问题。")
        else:
            logger.info("没有历史聊天记录，查询重写将仅基于当前问题。")

        QUERY_REWRITER_PROMPT = ChatPromptTemplate.from_messages(
            [
                ("system", """你是一个高级查询重写器，旨在帮助信息检索系统提高召回率。
                 根据用户提出的【问题】和{history_instruction}，生成 {num_queries} 条语义上相关但表述不同的检索子句。
                 这些子句应该：
                 1. 包含原始问题的核心意图。
                 2. {history_guidance}
                 3. 使用同义词、近义词或相关术语。
                 4. 考虑不同的关键词组合。
                 5. 如果问题涉及特定实体或概念，可以尝试从不同角度提问。
                 除了查询子句，不要输出任何其他内容，不要编号，不要解释，不要额外说明。
                 请确保每个子句都清晰、简洁，且旨在最大化检索的相关性。
                 """),
                MessagesPlaceholder("chat_history"),
                ("user", "【问题】: {question}\n输出：")
            ]
        )

        history_instruction = "【历史对话】" if use_history_for_rewrite else ""
        history_guidance = "考虑【历史对话】中可能存在的上下文和指代关系。" if use_history_for_rewrite else "忽略历史对话，聚焦当前问题。"

        try:
            rewriter_input = {
                "question": question,
                "chat_history": chat_history_list if use_history_for_rewrite else [],
                "num_queries": num_queries,
                "history_instruction": history_instruction,
                "history_guidance": history_guidance
            }

            raw_queries = (
                    QUERY_REWRITER_PROMPT
                    | self.llm
                    | StrOutputParser()
            ).invoke(rewriter_input)

            lines = []
            for line in raw_queries.splitlines():
                stripped_line = line.strip()
                if stripped_line and not stripped_line.startswith("【") and not stripped_line.startswith(
                        "---") and not stripped_line.startswith("###"):
                    lines.append(stripped_line)

            if question not in lines:
                lines.insert(0, question)

            unique_queries = list(dict.fromkeys(lines))
            logger.info(f"生成的唯一查询: {unique_queries[:num_queries]}")  # 记录最终查询
            return unique_queries[:num_queries]

        except Exception as e:
            logger.error(f"查询重写失败: {e}. 返回原始查询。", exc_info=True)
            return [question]

    def _multi_query_search(self, question: str, chat_history_list: List[Any], k_each: int = 8, max_total: int = 20,
                            history_relevance_threshold_for_rewrite: float = 0.5) -> List[Document]:

        queries = self._rewrite_queries(question, chat_history_list, num_queries=3,  # 减少查询数以提高速度
                                        history_relevance_threshold=history_relevance_threshold_for_rewrite)
        logger.info(f"用于检索的查询: {queries}")

        all_results_list = []

        for q in queries:
            # search 现在内部已经是 (BM25 + Vector) 的 RRF 结果了
            docs = self.search(q, k=k_each)
            if docs:
                all_results_list.append(docs)
        if not all_results_list:
            return []
        unique_docs = self._rrf_fusion(all_results_list, k=60)
        return unique_docs[:max_total]

    def search(self, query: str, k: int = 8) -> List[Document]:

        bm25_results = []
        chroma_results = []
        # 1. 获取 BM25 结果
        if self._bm25_retriever:
            try:
                # 动态设置 k
                self._bm25_retriever.k = k
                bm25_results = self._bm25_retriever.invoke(query)
                logger.debug(f"BM25 检索到 {len(bm25_results)} 个文档。")
            except Exception as e:
                logger.error(f"BM25 检索出错: {e}")

        # 2. 获取 Vector 结果
        if self._chroma_retriever:
            try:
                # 动态调整 Vector search 的 k
                chroma_results = self.vectorstore.similarity_search(query, k=k)
                logger.debug(f"Chroma 检索到 {len(chroma_results)} 个文档。")
            except Exception as e:
                logger.error(f"Chroma 检索出错: {e}")

        # 3. 如果只有一个检索器工作，直接返回结果
        if not bm25_results and not chroma_results:
            return []
        if not bm25_results:
            return chroma_results
        if not chroma_results:
            return bm25_results

        # 4. 执行 RRF 融合
        # 我们可以稍微调大 k 值以平滑排名，或者保持标准 60
        fused_docs = self._rrf_fusion([bm25_results, chroma_results], k=60)

        # 截断结果，返回前 k 个（或者稍微多一点供后续重排）
        return fused_docs[:k]

    def query(self, question: str, session_id: str = "default_session", k: int = 8, rerank_top_n: int = 7,
                  history_relevance_threshold_for_llm: float = 0.5,
                  history_relevance_threshold_for_rewrite: float = 0.5,
                  max_context_length: int = 4000) -> Dict[str, Any]:

            # 内部函数：获取历史记录对象
            def get_session_history(session_id: str) -> BaseChatMessageHistory:
                return SQLChatMessageHistory(session_id=session_id,
                                             connection=f"sqlite:///{self.chat_history_db_path}")

            # 内部函数：图片转Base64
            def encode_image(image_path):
                try:
                    with open(image_path, "rb") as image_file:
                        return base64.b64encode(image_file.read()).decode('utf-8')
                except Exception as e:
                    logger.error(f"图片编码失败 {image_path}: {e}")
                    return None


            confidence = 0.0
            relevant_docs_reranked = []  # 初始化，避免异常时引用报错

            try:
                # 1. 获取历史记录
                current_chat_history = get_session_history(session_id)
                current_chat_history_list = current_chat_history.messages

                # 2. 检索逻辑 (保持原逻辑：重写 -> 多路检索 -> RRF)
                relevant_docs = self._multi_query_search(
                    question,
                    current_chat_history_list,
                    k_each=k,
                    max_total=k * 2,
                    history_relevance_threshold_for_rewrite=history_relevance_threshold_for_rewrite
                )

                # 3. Rerank 和 文本上下文构建
                context_parts = []

                if relevant_docs:
                    logger.info(f"检索到 {len(relevant_docs)} 个文档，正在进行重新排名...")
                    if self.reranker:
                        sentence_pairs = [[question, doc.page_content] for doc in relevant_docs]
                        rerank_scores = self.reranker.predict(sentence_pairs)

                        # 组合并排序
                        doc_with_scores = sorted(zip(relevant_docs, rerank_scores), key=lambda x: x[1], reverse=True)

                        # 截取 Top N
                        relevant_docs_reranked = [doc for doc, score in doc_with_scores[:rerank_top_n]]

                        if relevant_docs_reranked:
                            confidence = doc_with_scores[0][1]  # 最高分作为基础置信度

                            # 构建文本 Context
                            current_context_length = 0
                            for doc in relevant_docs_reranked:
                                # 包含来源信息
                                source_info = f"\n--- Source: {doc.metadata.get('file_name', 'Unknown')}, Page: {doc.metadata.get('page_number', 'N/A')} ---\n"
                                chunk_content = doc.page_content + source_info

                                if current_context_length + len(chunk_content) <= max_context_length:
                                    context_parts.append(chunk_content)
                                    current_context_length += len(chunk_content)
                                else:
                                    remaining = max_context_length - current_context_length
                                    if remaining > len(source_info):
                                        context_parts.append(
                                            doc.page_content[:remaining - len(source_info)] + source_info)
                                    break

                            context = "\n\n".join(context_parts)
                        else:
                            context = ""
                            confidence = 0.2
                    else:
                        # 无 Reranker 的回退逻辑
                        relevant_docs_reranked = relevant_docs[:rerank_top_n]
                        context = "\n\n".join([d.page_content for d in relevant_docs_reranked])
                        confidence = 0.3
                else:
                    context = ""
                    confidence = 0.1

                # 4. 【新增】提取并处理相关图片
                relevant_images = []
                seen_images = set()

                if relevant_docs_reranked:
                    for doc in relevant_docs_reranked:
                        # 获取 image_paths，处理可能的类型差异 (list vs string)
                        img_paths_raw = doc.metadata.get("image_paths", [])
                        if isinstance(img_paths_raw, str):
                            try:
                                img_paths = ast.literal_eval(img_paths_raw)
                            except:
                                img_paths = []
                        else:
                            img_paths = img_paths_raw

                        # 验证并添加
                        if img_paths:
                            for img_path in img_paths:
                                if img_path and img_path not in seen_images and Path(img_path).exists():
                                    relevant_images.append(img_path)
                                    seen_images.add(img_path)

                # 限制图片数量，防止 Token 消耗过大 (例如最多3张)
                relevant_images = relevant_images[:3]
                if relevant_images:
                    logger.info(f"上下文包含 {len(relevant_images)} 张相关图片，将启用多模态回答。")

                # 5. 构建 System Prompt
                use_history_for_llm = False
                if current_chat_history_list:
                    relevance = self._evaluate_history_relevance(question, current_chat_history_list)
                    if relevance >= history_relevance_threshold_for_llm:
                        use_history_for_llm = True

                history_instruction = "结合【历史对话】" if use_history_for_llm else "忽略历史对话"

                system_prompt_text = f"""你是一个智能知识库助理。
    请根据以下【检索到的上下文】（可能包含文本、表格Markdown和图片）和{history_instruction}，准确、全面、详尽地回答用户的问题。

    "\n【重要原则】",
                "1.  **信息限定：** 你的回答必须严格限定在提供的【上下文】信息范围内。",
                "2.  **处理信息不足：** 如果上下文未能提供完整或直接的答案，请坦诚地指出“现有信息不足以完全回答此问题”或“在提供的资料中未找到直接答案。”。**严禁任何形式的臆测或补充外部信息。**",
                "3.  **整合与提炼：** 当上下文包含多个相关片段时，请将它们整合起来，形成一个连贯、有逻辑的回答。**如果信息足够，请尽可能详细地展开。**",
                # 强调详细展开
                "4.  **结构化呈现：** 如果问题复杂，可以考虑使用列表、分点说明等方式使答案更易读。",
                "5.  **语言风格：** 保持专业、客观和中立的语气。"
    """

                # 6. 构建 User Message (多模态)
                content_blocks = []

                # 6.1 添加文本部分
                text_payload = f"【上下文】\n{context}\n\n【用户问题】\n{question}\n\n【你的详细回答】"
                content_blocks.append({"type": "text", "text": text_payload})

                # 6.2 添加图片部分
                for img_path in relevant_images:
                    base64_img = encode_image(img_path)
                    if base64_img:
                        content_blocks.append({
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_img}",
                                "detail": "auto"
                            }
                        })

                # 7. 组装最终消息列表 (System + History + User)
                final_messages = [SystemMessage(content=system_prompt_text)]

                if use_history_for_llm:
                    final_messages.extend(current_chat_history_list)

                final_messages.append(HumanMessage(content=content_blocks))

                # 8. 调用 LLM 生成回答
                # 注意：这里直接 invoke LLM，绕过了 RunnableWithMessageHistory，因为多模态消息构造比较特殊
                response = self.llm.invoke(final_messages)
                answer = response.content

                # 9. 手动保存历史记录 (存纯文本即可)
                # 用户问题存文本
                current_chat_history.add_user_message(question)
                # AI 回答存文本
                current_chat_history.add_ai_message(answer)

                # 10. LLM 自评估 (仅基于文本上下文和回答，节省图片Token)
                # 这里的逻辑保持不变，用于生成 confidence
                llm_eval_prompt = ChatPromptTemplate.from_messages([
                ("system", """你是一个评估助手。你的任务是根据提供的原始【问题】、【检索到的上下文】和【AI生成的答案】，判断AI答案的质量和忠实性。
                            请严格根据以下三点进行评估：
                            1. 答案是否完全基于提供的上下文？(Faithfulness) - 如果答案包含上下文之外的信息，则忠实度低。
                            2. 答案是否直接、完整地回答了问题？(Relevance & Completeness) - 答案是否涵盖了问题的所有方面，并且与问题高度相关。
                            3. 综合上述，请给出一个 0 到 1 之间的整体置信度分数，表示你对这个AI答案的信任程度。

                            **重要提示：请只输出一个介于0.0到1.0之间的浮点数，不要包含任何其他文字、解释或标点符号。**
                            """),  # 增加更严格的提示
                ("user", "【问题】: {question}\n【上下文】: {context}\n【AI生成的答案】: {answer}\n最终置信度分数:")
            ])

                eval_chain = llm_eval_prompt | self.llm | StrOutputParser()

                try:
                    # 评估时不传图片，只传文本上下文
                    eval_res = eval_chain.invoke({"question": question, "context": context, "answer": answer})
                    import re
                    match = re.search(r"(\d+\.\d+)", eval_res)
                    if match:
                        llm_conf = float(match.group(1))
                        llm_conf = max(0.0, min(1.0, llm_conf))
                        # 综合评分：LLM评估占70%，检索Rerank分占30%
                        confidence = (llm_conf * 0.7 + confidence * 0.3) if relevant_docs else llm_conf
                except Exception as e:
                    logger.warning(f"自评估失败: {e}，使用默认置信度")

                # 11. 整理 Sources 返回结果
                sources = []
                seen_source_identifiers = set()

                # 添加文档来源
                if relevant_docs_reranked:
                    # 重新找到对应的 rerank score (如果有的话)
                    doc_score_map = {doc.metadata.get('chunk_id'): score
                                     for doc, score in zip([d for d, s in doc_with_scores], [s for d, s in
                                                                                             doc_with_scores])} if 'doc_with_scores' in locals() else {}

                    for doc in relevant_docs_reranked:
                        file_name = doc.metadata.get('file_name', 'Unknown')
                        page_number = doc.metadata.get('page_number', 'N/A')
                        chunk_id = doc.metadata.get('chunk_id')
                        score = doc_score_map.get(chunk_id, 0.0)

                        identifier = f"{file_name}#page_{page_number}"
                        if identifier not in seen_source_identifiers:
                            sources.append({
                                "type": "text",
                                "file": file_name,
                                "page": page_number,
                                "rerank_score": score,
                                "content_preview": doc.page_content[:100] + "..."
                            })
                            seen_source_identifiers.add(identifier)

                # 添加图片来源信息
                for img_path in relevant_images:
                    sources.append({
                        "type": "image",
                        "file": Path(img_path).name,  # 图片文件名
                        "page": "Image Reference",
                        "rerank_score": 1.0,  # 图片作为直接证据，置信度设为高
                        "content_preview": f"Image Path: {img_path}"
                    })

                return {
                    "answer": answer,
                    "sources": sources,
                    "confidence": confidence,
                }

            except Exception as e:
                logger.exception(f"查询时出错: {e}")
                return {
                    "answer": f"抱歉, 处理您的查询时出现了错误: {str(e)}",
                    "sources": [],
                    "confidence": 0.0,
                }
    def clear_chat_history(self, session_id: Optional[str] = None):
        if session_id:
            try:
                history = SQLChatMessageHistory(session_id=session_id,
                                                connection=f"sqlite:///{self.chat_history_db_path}")
                history.clear()
                logger.info(f"会话 '{session_id}' 的对话历史已从数据库中清空。")
            except Exception as e:
                logger.warning(f"清空会话 '{session_id}' 历史时出错: {e}", exc_info=True)
        else:
            if Path(self.chat_history_db_path).exists():
                try:
                    os.remove(self.chat_history_db_path)
                    logger.info(f"所有会话的对话历史数据库文件 '{self.chat_history_db_path}' 已删除。")
                    Path(self.chat_history_db_path).parent.mkdir(parents=True, exist_ok=True)
                except Exception as e:
                    logger.error(f"删除所有会话历史数据库文件时出错: {e}", exc_info=True)
            else:
                logger.info("聊天历史数据库文件不存在，无需清空。")

    def list_documents(self) -> Dict[str, Any]:
        total_chunks_in_memory = len(self.documents)
        chroma_index_size = 0

        self._get_or_create_chroma()
        try:
            if self.vectorstore is not None:
                chroma_index_size = self.vectorstore._collection.count()
            else:
                chroma_index_size = 0
        except Exception as e:
            logger.error(f"获取 Chroma 集合实体数量时出错: {e}", exc_info=True)
            chroma_index_size = -1

        return {
            "total_chunks_in_memory": total_chunks_in_memory,
            "vector_index_size": chroma_index_size,
            "documents": self.doc_metadata,
        }

    def delete_document(self, file_path: str):
        file_path_obj = Path(file_path).resolve()
        file_path_str = str(file_path_obj)

        if file_path_str in self.doc_metadata:
            logger.info(f"正在删除文档: {file_path_str}")
            self._delete_file_from_vectorstore(file_path_str)

            self.chunk_store.delete_file(file_path_str)
            self.documents = [doc for doc in self.documents if doc.metadata.get('source') != file_path_str]

            del self.doc_metadata[file_path_str]
            del self._file_hashes[file_path_str]

            self._save_knowledge_base()
            self._initialize_retrievers()
            logger.info(f"已删除文档: {file_path_str}")
        else:
            logger.warning(f"文档不存在于知识库元数据中，无法删除: {file_path_str}")
            self._delete_file_from_vectorstore(file_path_str)

    def delete_all_documents(self):
        logger.info("正在删除所有文档并清空知识库...")
        self.documents = []
        self.doc_metadata = {}
        self._file_hashes = {}
        self.chunk_store.clear()
        self.clear_chat_history(session_id=None)

        if Path(self.chroma_persist_path).exists():
            try:
                shutil.rmtree(self.chroma_persist_path)
                logger.info(f"已删除 Chroma 数据库目录: {self.chroma_persist_path}")
                Path(self.chroma_persist_path).mkdir(parents=True, exist_ok=True)
            except Exception as e:
                logger.error(f"删除 Chroma 数据库目录时出错: {e}", exc_info=True)
        else:
            logger.info("Chroma 数据库目录不存在，无需删除。")

        self.vectorstore = None

        self._save_knowledge_base()
        self._initialize_retrievers()
        logger.info("所有文档已删除，知识库已清空。")

    def list_all_session_ids(self) -> List[str]:
        try:
            conn_str = f"sqlite:///{self.chat_history_db_path}"
            temp_history = SQLChatMessageHistory(session_id="dummy", connection=conn_str)

            # 从 temp_history 获取 session_maker
            Session = temp_history.session_maker

            with Session() as session:
                from sqlalchemy import text
                result = session.execute(text("SELECT DISTINCT session_id FROM langchain_chat_histories")).fetchall()
                session_ids = [row[0] for row in result]
                return session_ids
        except Exception as e:
            logger.error(f"获取所有会话ID时出错: {e}", exc_info=True)
            return []

'''
def main():
    deepseek_api_key = os.getenv("DEEPSEEK_API_KEY")
    if not deepseek_api_key:
        logger.error("程序退出: 请先在环境变量中设置 OPENAI_API_KEY。")
        print("请在运行程序前设置 OPENAI_API_KEY 环境变量，例如：")
        print("export OPENAI_API_KEY='你的API密钥'   （Linux/macOS）")
        print("set OPENAI_API_KEY=你的API密钥         （Windows CMD）")
        print("$env:OPENAI_API_KEY='你的API密钥'      （PowerShell）")
        return 1

    try:
        kb = MultiDocumentKnowledgeBase("./turbine_machine", deepseek_api_key=deepseek_api_key,llm_instance=None)

        current_session_id = "user1_session"

        while True:
            print("\n" + "=" * 60)
            print(f"多文档知识库系统 (当前会话ID: {current_session_id})")
            print("=" * 60)
            print("1. 添加单个文档")
            print("2. 从文件夹添加所有文档")
            print("3. 查询问题 (支持历史对话)")
            print("4. 列出所有文档")
            print("5. 删除文档")
            print("6. 删除所有文档")
            print("7. 清空当前会话历史")
            print("8. 切换/创建会话")
            print("9. 列出所有活跃会话ID")
            print("10. 清空所有会话历史")
            print("11. 退出")
            print("=" * 60)

            choice = input("请输入您的选择 (1-11): ").strip()

            if choice == '1':
                file_path = input("请输入要添加的文档路径: ").strip()
                try:
                    kb.add_document(file_path)
                    print(f"文档 '{file_path}' 添加成功。")
                except FileNotFoundError:
                    print("错误: 文件不存在。")
                except Exception as e:
                    print(f"添加文档时发生错误: {e}")

            elif choice == '2':
                directory_path = input("请输入要添加文档的文件夹路径: ").strip()
                try:
                    kb.add_documents_from_directory(directory_path)
                    print("文件夹内所有支持的文档已处理完成!")
                except Exception as e:
                    print(f"处理文件夹时失败: {e}")

            elif choice == "3":
                question = input("请输入您的问题: ").strip()
                if question:
                    # 允许用户输入更多参数
                    k_val = int(input("请输入每个检索器检索的文档数量 k (默认8): ") or 8)
                    rerank_top_n_val = int(input("请输入Rerank后保留的顶部文档数量 (默认5): ") or 5)
                    history_relevance_threshold_for_llm = float(
                        input("请输入LLM使用历史对话的相关度阈值 (0.0-1.0, 默认0.5): ") or 0.5)
                    history_relevance_threshold_for_rewrite = float(
                        input("请输入查询重写使用历史对话的相关度阈值 (0.0-1.0, 默认0.5): ") or 0.5)
                    max_context_length_val = int(input("请输入传递给LLM的最大上下文长度 (默认4000): ") or 4000)

                    print("\n正在查询...")
                    result = kb.query(question, session_id=current_session_id,
                                      k=k_val,
                                      rerank_top_n=rerank_top_n_val,
                                      history_relevance_threshold_for_llm=history_relevance_threshold_for_llm,
                                      history_relevance_threshold_for_rewrite=history_relevance_threshold_for_rewrite,
                                      max_context_length=max_context_length_val)
                    print(f"\n回答: {result['answer']}")
                    print("\n" + "=" * 20 + " 相关资料 " + "=" * 20)
                    if result["sources"]:
                        for source in result["sources"]:
                            print(
                                f"  - 文件: {source['file']}, 页面: {source['page']}, Rerank分数: {source['rerank_score']:.4f}")
                            # print(f"    内容预览: {source['content_preview']}") # 如果需要，可以打印预览
                    else:
                        print("  无相关资料。")
                    print("=" * 20 + " 结束 " + "=" * 20)
                    print(f"最终置信度: {result['confidence']:.4f}")

            elif choice == "4":
                documents_info = kb.list_documents()
                print("\n" + "=" * 50)
                print("知识库文档列表")
                print("=" * 50)
                print(f"内存中文档片段总数 (BM25使用): {documents_info['total_chunks_in_memory']}")
                if documents_info['vector_index_size'] == -1:
                    print(f"向量索引实体总数: 无法精确获取 (Chroma可能需要手动检查或更新)")
                else:
                    print(f"向量索引实体总数: {documents_info['vector_index_size']}")
                print(f"已添加到知识库的文件:")
                if documents_info["documents"]:
                    for file_path, metadata in documents_info["documents"].items():
                        print(f"  - 文件: {Path(file_path).name}")
                        print(f"    路径: {file_path}")
                        print(f"    类型: {metadata['doc_type']}")
                        print(f"    片段数: {metadata['chunk_count']}")
                else:
                    print("  知识库中没有文档。")
                print("=" * 50)

            elif choice == "5":
                file_path = input("请输入要删除的文档的完整路径: ").strip().strip('"')
                if file_path:
                    try:
                        kb.delete_document(file_path)
                        print("文档删除成功!")
                    except Exception as e:
                        print(f"删除文档失败: {e}")
                else:
                    print("文件路径不能为空。")

            elif choice == "6":
                confirm = input("确定要删除所有文档吗？此操作不可逆，请输入 'yes' 确认: ").strip().lower()
                if confirm == 'yes':
                    try:
                        kb.delete_all_documents()
                        print("所有文档已成功删除。")
                    except Exception as e:
                        print(f"删除所有文档失败: {e}")
                else:
                    print("操作已取消。")

            elif choice == "7":
                kb.clear_chat_history(session_id=current_session_id)
                print(f"当前会话 '{current_session_id}' 的对话历史已清空。")

            elif choice == "8":
                new_session_id = input(f"请输入新的会话ID (当前: {current_session_id}): ").strip()
                if new_session_id:
                    current_session_id = new_session_id
                    print(f"已切换到会话ID: {current_session_id}")
                else:
                    print("会话ID不能为空，未切换。")

            elif choice == "9":
                print("\n" + "=" * 30)
                print("所有活跃会话ID (从数据库加载)")
                print("=" * 30)
                session_ids = kb.list_all_session_ids()
                if session_ids:
                    for sid in session_ids:
                        print(f"- {sid}")
                else:
                    print("无活跃会话历史。")
                print("=" * 30)

            elif choice == "10":
                confirm = input("确定要清空所有会话历史吗？此操作不可逆，请输入 'yes' 确认: ").strip().lower()
                if confirm == 'yes':
                    kb.clear_chat_history(session_id=None)
                    print("所有会话的历史已清空。")
                else:
                    print("操作已取消。")

            elif choice == "11":
                print("感谢使用!")
                break

            else:
                print("无效选择,请重新输入。")

    except Exception as e:
        logger.exception(f"程序运行出错: {e}")
        print(f"程序出现错误: {e}")


if __name__ == "__main__":
    exit_code = main()
    if exit_code:
        exit(exit_code)
'''
